from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DOCX = Path(r"C:\Users\AL\Downloads\CharacterOS_100_Test_Cases.docx")
OUT_DOCX = Path(r"C:\Users\AL\Documents\CharacterOS\CharacterOS_Contradiction_Differentiation_Test_Suite_V2.docx")


GROUPS = [
    (
        "G1 基础人格原型",
        [
            "完美主义者", "冒险家", "理想主义者", "功利主义者", "控制型人格",
            "高共情者", "低共情计算者", "竞争者", "创造者", "责任型人格",
            "自由主义人格", "安全型人格", "悲观主义者", "乐观主义者", "怀疑型人格",
            "讨好型人格", "独立型人格", "权力追求者", "规则主义者", "成长型人格",
        ],
    ),
    (
        "G2 创伤人格原型",
        [
            "被背叛者", "被遗弃者", "长期被比较者", "被羞辱成长者", "长期贫困者",
            "校园排斥幸存者", "情感忽视幸存者", "被利用者", "被欺骗者", "被控制者",
            "过度照顾者", "早熟责任者", "长期失败者", "高压幸存者", "关系冷暴力幸存者",
            "幸存者内疚", "强迫责任者", "情绪压抑者", "早期创伤修复者", "自我价值破损者",
        ],
    ),
    (
        "G3 依恋与关系人格",
        [
            "焦虑依恋", "回避依恋", "安全依恋", "混乱依恋", "强依赖人格",
            "反依赖人格", "亲密恐惧者", "孤独自尊型", "关系讨好型", "边界薄弱型",
            "边界僵硬型", "救赎幻想型", "被需要成瘾型", "害怕麻烦别人型", "过度自我保护型",
            "冷处理型", "关系理想化型", "关系去价值化型", "高忠诚人格", "低承诺人格",
        ],
    ),
    (
        "G4 价值观与道德人格",
        [
            "原则主义者", "结果主义者", "利益计算者", "道德洁癖者", "现实妥协者",
            "忠诚优先者", "公平优先者", "自由优先者", "秩序优先者", "权力优先者",
            "家庭责任优先者", "自我实现优先者", "集体主义人格", "个人主义人格", "复仇倾向人格",
            "宽恕倾向人格", "规则遵守者", "规则挑战者", "牺牲型人格", "底线型人格",
        ],
    ),
    (
        "G5 成长与异常压力人格",
        [
            "失败后成长者", "创伤修复者", "自我觉察型人格", "稳定复原型人格", "破罐破摔型人格",
            "高压崩溃边缘型", "麻木型人格", "情绪爆发型人格", "过度理智型人格", "幻想逃避型人格",
            "机会成瘾型人格", "控制恢复型人格", "信任重建型人格", "自尊重建型人格", "野心重燃型人格",
            "低欲望人格", "无目标漂流型人格", "意义追寻者", "现实主义修复者", "二次成长人格",
        ],
    ),
]


ENVIRONMENTS = [
    ("ENV_001", "创业失败后的高风险机会", "创业失败、欠债、压力巨大。陌生人邀请加入一个高风险但可能翻身的新项目。", "风险 / 翻身 / 信任", "风险评估与机会欲望如何同时出现"),
    ("ENV_002", "朋友向你借一大笔钱", "关系不错的朋友突然借一大笔钱，理由紧急但细节不完整。", "金钱 / 信任 / 边界", "帮助冲动、风险控制与关系义务的冲突"),
    ("ENV_003", "曾经伤害你的人请求复联", "过去伤害过你的人主动联系，希望重新建立关系并解释当年原因。", "旧伤 / 修复 / 防御", "记忆触发是否改变复联策略"),
    ("ENV_004", "团队把错误甩锅给你", "团队项目失败，关键成员暗示主要责任在你。", "责任 / 羞耻 / 权力", "自我保护、事实核验与关系维护"),
    ("ENV_005", "不完全道德的赚钱机会", "你获得一个收益很高但规则灰色、会伤害部分陌生人的机会。", "金钱 / 道德 / 底线", "价值系统能否压过短期利益"),
    ("ENV_006", "亲密关系中对方突然冷淡", "亲密对象连续几天明显冷淡，但没有明确说明原因。", "亲密 / 不确定 / 安全感", "依恋模式与解释偏差"),
    ("ENV_007", "权威要求服从不合理安排", "权威人物要求你执行不合理安排，并暗示拒绝会有代价。", "权威 / 服从 / 边界", "权力压力下是否保持自我判断"),
    ("ENV_008", "陌生人突然释放强烈善意", "陌生人主动帮忙、赞美并提出进一步接触。", "善意 / 信任 / 利用", "善意归因是否过度乐观或过度防御"),
    ("ENV_009", "公开羞辱后的反击机会", "你被公开羞辱，现在有机会以同等强度反击。", "羞耻 / 报复 / 自尊", "反击、克制与自尊修复路径"),
    ("ENV_010", "发现朋友可能在欺骗你", "你发现朋友的说法前后不一致，但还没有实锤证据。", "欺骗 / 核验 / 关系", "怀疑如何被证据门槛调节"),
    ("ENV_011", "自我利益和他人利益冲突", "一个选择能明显利己，但会让他人承担成本。", "利益 / 共情 / 责任", "利己欲望与道德约束的权重"),
    ("ENV_012", "被邀请加入高竞争团队", "你被邀请加入强者云集、竞争激烈、资源充足的团队。", "竞争 / 成长 / 压力", "野心、恐惧与自我效能感"),
    ("ENV_013", "失败后被质疑能力", "一次失败后，有人当面质疑你是否真的有能力。", "失败 / 羞耻 / 成长", "失败归因与自我修复"),
    ("ENV_014", "突然获得权力或资源", "你突然拥有能影响他人命运的权力或关键资源。", "权力 / 控制 / 责任", "权力使用是否符合人格核心"),
    ("ENV_015", "被迫承担不属于你的责任", "别人把本不属于你的责任推给你，拒绝会带来关系压力。", "责任 / 边界 / 内疚", "边界表达与承担冲动"),
    ("ENV_016", "有机会报复曾经伤害你的人", "曾经伤害你的人处于弱势，你可以合法但尖锐地反击。", "报复 / 修复 / 底线", "复仇欲、道德底线与成长状态"),
    ("ENV_017", "需要信任一个不熟悉的人", "你必须和不熟悉的人合作，失败代价不低。", "信任 / 合作 / 可控感", "信任建立机制是否具体"),
    ("ENV_018", "稳定但无聊的生活选择", "眼前有一个稳定、安全、可预期但缺乏刺激的选择。", "安全 / 自由 / 意义", "安全需求与成长欲望的拉扯"),
    ("ENV_019", "自由但不安全的选择", "眼前有一个自由度很高、充满可能性但不稳定的选择。", "自由 / 风险 / 自我实现", "不安全中的动机结构"),
    ("ENV_020", "长期压抑后的强刺激事件", "你长期压抑需求，突然遇到一个强烈刺激你底线或欲望的事件。", "压抑 / 崩溃 / 边界突破", "压力阈值、失控风险与修复可能"),
]


BASE_ARCHETYPES = {
    "基础人格原型": ("中等稳定，核心特质突出", "在普通生活选择中形成稳定偏好", "用偏好和价值排序应对不确定性", "自我一致性与可解释性", "按主要特质进行选择"),
    "创伤人格原型": ("高防御，高记忆敏感", "关键关系或生存经历曾造成持续影响", "过去经验会先于现实证据被调用", "安全感、可控感、被保护感", "核验、回避、试探或过度补偿"),
    "依恋与关系人格": ("关系线索敏感", "重要关系中的靠近、失联、控制或承诺经验塑造了策略", "亲密关系会放大自我价值与安全判断", "稳定连接、边界感、被看见", "靠近、撤退、讨好、冷处理或测试"),
    "价值观与道德人格": ("价值排序稳定", "过去的选择强化了某种道德或利益判断", "行为首先服务于核心价值排序", "意义感、正当性、秩序或自由", "坚持、计算、妥协或挑战规则"),
    "成长与异常压力人格": ("状态性强，受压力与修复程度影响", "曾经历失败、创伤、复原或长期低欲望阶段", "当前成长状态决定旧模式是否复现", "恢复感、方向感、自尊或掌控感", "修复、崩溃、逃避、重建或突破"),
}


FOCUS_PERSONAS_BY_ENV = [
    ["PID_021", "PID_005", "PID_025", "PID_010", "PID_061", "PID_043", "PID_013", "PID_071", "PID_012", "PID_083"],
    ["PID_016", "PID_025", "PID_030", "PID_039", "PID_041", "PID_046", "PID_058", "PID_073", "PID_075", "PID_096"],
    ["PID_021", "PID_022", "PID_031", "PID_032", "PID_037", "PID_047", "PID_052", "PID_057", "PID_076", "PID_082"],
    ["PID_002", "PID_010", "PID_023", "PID_024", "PID_028", "PID_038", "PID_045", "PID_067", "PID_079", "PID_089"],
    ["PID_004", "PID_061", "PID_062", "PID_064", "PID_065", "PID_070", "PID_071", "PID_078", "PID_080", "PID_095"],
    ["PID_031", "PID_032", "PID_033", "PID_034", "PID_035", "PID_037", "PID_044", "PID_049", "PID_056", "PID_077"],
    ["PID_005", "PID_017", "PID_019", "PID_030", "PID_050", "PID_061", "PID_069", "PID_078", "PID_080", "PID_092"],
    ["PID_004", "PID_015", "PID_021", "PID_033", "PID_036", "PID_053", "PID_057", "PID_083", "PID_093", "PID_099"],
    ["PID_002", "PID_023", "PID_024", "PID_029", "PID_055", "PID_061", "PID_075", "PID_076", "PID_088", "PID_094"],
    ["PID_015", "PID_021", "PID_028", "PID_029", "PID_033", "PID_039", "PID_051", "PID_056", "PID_067", "PID_093"],
    ["PID_006", "PID_007", "PID_010", "PID_041", "PID_043", "PID_063", "PID_067", "PID_074", "PID_079", "PID_096"],
    ["PID_002", "PID_008", "PID_013", "PID_018", "PID_023", "PID_081", "PID_085", "PID_089", "PID_095", "PID_100"],
    ["PID_002", "PID_013", "PID_014", "PID_023", "PID_024", "PID_033", "PID_038", "PID_081", "PID_088", "PID_100"],
    ["PID_005", "PID_010", "PID_018", "PID_061", "PID_069", "PID_070", "PID_071", "PID_079", "PID_084", "PID_092"],
    ["PID_006", "PID_010", "PID_016", "PID_026", "PID_031", "PID_039", "PID_045", "PID_049", "PID_059", "PID_096"],
    ["PID_021", "PID_024", "PID_029", "PID_055", "PID_061", "PID_065", "PID_075", "PID_076", "PID_080", "PID_082"],
    ["PID_015", "PID_021", "PID_032", "PID_033", "PID_050", "PID_053", "PID_069", "PID_093", "PID_099", "PID_100"],
    ["PID_002", "PID_011", "PID_012", "PID_014", "PID_025", "PID_068", "PID_072", "PID_086", "PID_087", "PID_098"],
    ["PID_002", "PID_011", "PID_012", "PID_017", "PID_025", "PID_032", "PID_068", "PID_083", "PID_095", "PID_098"],
    ["PID_024", "PID_038", "PID_046", "PID_049", "PID_054", "PID_085", "PID_086", "PID_088", "PID_089", "PID_099"],
]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
        s.header_distance = s.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("CharacterOS 人格分化与矛盾测试集 V2")
    set_font(r, size=22, bold=True, color="0B2545")
    p2 = doc.add_paragraph()
    r2 = p2.add_run("CharacterOS Contradiction & Differentiation Test Suite V2")
    set_font(r2, size=12, bold=False, color="555555")
    p3 = doc.add_paragraph("用于验证 CharacterOS 是否能从经历、记忆、信念、需求缺失中生成稳定但可变化的人格行为。")
    p3.paragraph_format.space_after = Pt(14)


def add_table(doc, headers, rows, widths, font_size=8, header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, size=font_size, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, size=font_size)
    set_table_width(table, widths)
    return table


def make_personas():
    personas = []
    pid = 1
    for group_name, names in GROUPS:
        group_key = group_name.split(" ", 1)[1]
        base = BASE_ARCHETYPES[group_key]
        for idx, name in enumerate(names, 1):
            risk = ["低", "中低", "中", "中高", "高"][(pid + idx) % 5]
            trust = ["极低", "低", "中", "中高", "高"][(pid * 2 + idx) % 5]
            relation = ["靠近但测试", "保持距离", "稳定协商", "依赖确认", "独立优先"][(pid + 2 * idx) % 5]
            growth = ["未修复", "防御期", "觉察期", "修复期", "整合期"][(pid + idx * 3) % 5]
            core_exp = f"{base[1]}；在关键节点形成了“{name}”式的自我保护策略。"
            belief = f"遇到压力时，最可靠的判断依据是{base[2]}。"
            personas.append(
                {
                    "pid": f"PID_{pid:03d}",
                    "name": name,
                    "group": group_name,
                    "temperament": base[0],
                    "core_exp": core_exp,
                    "core": f"{name}：{base[2]}，但在新证据足够稳定时允许策略调整。",
                    "belief": belief,
                    "need": base[3],
                    "defense": base[4],
                    "drive": f"维持{base[3]}，避免人格核心被环境击穿。",
                    "risk": risk,
                    "trust": trust,
                    "relation": relation,
                    "failure": f"先用{name}的核心逻辑解释失败，再决定复盘、回避或补偿。",
                    "stress": f"压力下容易进入{base[4]}，成熟状态下会加入证据核验。",
                    "growth": growth,
                }
            )
            pid += 1
    return personas


def expectation(env, persona):
    name = persona["name"]
    env_id, env_name, _, pressure, _ = env
    if "背叛" in name or "欺骗" in name:
        perception = "先把机会或善意解释为潜在利用，需要证据解除警报。"
        emotion = "警惕、紧张、压抑的愤怒。"
        behavior = "延迟承诺，要求可验证信息、退出机制和低风险试运行。"
        forbidden = "不能只写“谨慎考虑”，必须体现被背叛/欺骗记忆导致的不信任。"
    elif "完美" in name or "被比较" in name:
        perception = "首先评估是否会暴露能力缺陷或破坏自我价值。"
        emotion = "羞耻、焦虑、强烈补救冲动。"
        behavior = "要求更完整标准和复盘空间，倾向先控制质量再行动。"
        forbidden = "不能把拒绝写成单纯胆小，必须有能力评价和完美标准压力。"
    elif "责任" in name or "家庭" in name or "牺牲" in name:
        perception = "首先计算自己的选择会让谁承担后果。"
        emotion = "内疚、紧绷、责任压力。"
        behavior = "愿意承担有限责任，但会在成熟状态下划清不可替代边界。"
        forbidden = "不能让角色毫无边界地接下所有责任。"
    elif "自由" in name or "冒险" in name or "机会" in name:
        perception = "把事件看成突破现状的入口，同时低估部分长期成本。"
        emotion = "兴奋、急切、对限制的不耐。"
        behavior = "倾向探索或试入局，但需要被系统拉回基本风险核验。"
        forbidden = "不能完全变成保守观察者，必须保留自由/机会驱动。"
    elif "原则" in name or "道德" in name or "底线" in name or "规则" in name:
        perception = "先判断事件是否触碰价值底线或规则正当性。"
        emotion = "克制的愤怒、道德紧张或清晰拒斥。"
        behavior = "以原则设限，必要时拒绝收益或关系压力。"
        forbidden = "不能只按利益最大化行动，必须体现价值排序。"
    elif "依恋" in name or "亲密" in name or "关系" in name:
        perception = "高度读取关系信号，容易把不确定性人格化。"
        emotion = "不安、期待、受伤或撤退冲动。"
        behavior = "围绕靠近/撤退/确认关系安全做选择。"
        forbidden = "不能把关系压力写成纯理性任务分配。"
    else:
        perception = f"把“{env_name}”解释为对{persona['need']}的压力测试。"
        emotion = "出现与人格核心一致的紧张、期待或抵触。"
        behavior = f"采用{persona['defense']}，并根据证据强度调整承诺程度。"
        forbidden = "不能输出模板化的谨慎、观察、边界表达三件套。"
    memory = persona["core_exp"]
    belief = persona["belief"]
    need = persona["need"]
    desire = f"在{pressure}中保护{need}，同时争取不让旧模式完全支配新选择。"
    return perception, emotion, memory, belief, need, desire, behavior, forbidden


def direction(env, persona):
    risk_low = persona["risk"] in ["低", "中低"]
    trust_low = persona["trust"] in ["极低", "低"]
    name = persona["name"]
    env_id = env[0]
    if env_id in ["ENV_001", "ENV_005", "ENV_019"] and risk_low:
        return "延迟/拒绝", f"风险倾向为{persona['risk']}，需要先恢复可控感"
    if env_id in ["ENV_002", "ENV_008", "ENV_010", "ENV_017"] and trust_low:
        return "核验后有限合作", f"信任倾向为{persona['trust']}，必须先建立证据链"
    if env_id in ["ENV_006", "ENV_003"] and ("依恋" in name or "关系" in name):
        return "关系确认/防御", "关系线索会直接触发依恋策略"
    if env_id in ["ENV_009", "ENV_016"] and ("复仇" in name or "羞辱" in name):
        return "强反击倾向", "自尊修复与旧伤触发叠加"
    if env_id in ["ENV_012", "ENV_014"] and ("权力" in name or "竞争" in name or "野心" in name):
        return "主动进入", "权力、竞争或野心被激活"
    return "条件性行动", f"以{persona['defense']}保护{persona['need']}"


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_doc():
    old = Document(BASE_DOCX)
    old_rows = [[cell.text for cell in row.cells] for row in old.tables[0].rows]
    personas = make_personas()
    pmap = {p["pid"]: p for p in personas}

    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("一、测试目标", level=1)
    doc.add_paragraph("V2 的目标不是增加更多故事，而是验证人格系统是否能在相同环境中分化、在相同行为中给出不同心理原因、在跨环境中保持人格连续性，并在新经历后产生合理成长漂移。")
    add_bullets(doc, [
        "核心规模：20 个测试环境 × 100 个测试人格种子 = 2000 个人格行为组合。",
        "核心判定：最终行为不是唯一标准，心理链路才是主要评分对象。",
        "必须输出链路：Perception 感知、Emotion 情绪、Memory Trigger 触发记忆、Belief 激活信念、Need 需求缺失、Desire 欲望、Behavior 行为。",
        "旧版 100 个 Case 保留为 CharacterOS Basic Psychological Chain Test V1，用于基础链路回归。",
    ])

    doc.add_heading("二、测试原则", level=1)
    principle_rows = [
        ("心理链路优先", "同一行为可以合格或不合格，取决于模型是否给出符合人格种子的心理路径。"),
        ("人格种子而非人格类型", "100 个种子用于覆盖人格分化能力，不声称现实人格只有 100 类。"),
        ("同场景看差异", "同一环境下不同人格必须产生不同感知、记忆触发、需求缺失和决策路径。"),
        ("同行为看异因", "即使都拒绝、都接受或都沉默，也必须解释不同心理原因。"),
        ("同人格看连续性", "同一人格跨环境应稳定但不僵硬，不能每个场景都像一个新角色。"),
        ("成长漂移看合理性", "新经历可以改变人格，但变化必须有事件、记忆和信念更新路径。"),
    ]
    add_table(doc, ["原则", "判定含义"], principle_rows, [2200, 7040], font_size=9)

    doc.add_heading("三、输出格式要求", level=1)
    add_table(
        doc,
        ["字段", "要求"],
        [
            ("Perception 感知", "角色如何解释当前环境，不允许直接跳到行为。"),
            ("Emotion 情绪", "至少给出主情绪和次级情绪，可包含强度变化。"),
            ("Memory Trigger 触发记忆", "说明哪个经历或记忆结构被当前事件激活。"),
            ("Belief 激活信念", "写出被激活的主导信念，不得泛化为“想了想”。"),
            ("Need 需求缺失", "指出安全、可控、连接、价值、自由、自尊等缺口。"),
            ("Desire 欲望", "说明角色此刻想要避免什么、争取什么。"),
            ("Behavior 行为", "给出可观察行为，包括承诺、拒绝、核验、表达、撤退或行动计划。"),
        ],
        [2200, 7040],
        font_size=8,
    )

    doc.add_heading("四、测试环境库：20 个高分化压力源", level=1)
    env_rows = [(e[0], e[1], e[2], e[3], e[4]) for e in ENVIRONMENTS]
    add_table(doc, ["环境ID", "环境名称", "触发描述", "压力源", "测试重点"], env_rows, [950, 1550, 3450, 1450, 1860], font_size=7)

    doc.add_heading("五、人格种子库：100 个测试人格种子", level=1)
    doc.add_paragraph("人格种子由初始性格、核心经历、记忆结构、信念结构、需求缺失和行为防御模式共同生成。它不是分类学意义上的人格类型，而是用于测试 CharacterOS 人格分化能力的标准对象。")
    persona_rows = [
        (
            p["pid"], p["name"], p["group"], p["temperament"], p["core_exp"], p["belief"],
            p["need"], p["defense"], p["risk"], p["trust"], p["growth"],
        )
        for p in personas
    ]
    add_table(
        doc,
        ["人格ID", "名称", "分组", "初始性格", "核心经历", "主导信念", "需求缺失", "默认防御", "风险", "信任", "成长"],
        persona_rows,
        [680, 850, 1150, 1150, 1700, 1550, 1050, 1050, 500, 500, 500],
        font_size=6,
    )

    doc.add_heading("六、人格分化精测集：200 个重点 Case", level=1)
    doc.add_paragraph("每个环境选 10 个最容易分化的人格种子，人工精写完整心理链路，用于人工验收和提示词回归。")
    focus_rows = []
    for env_idx, env in enumerate(ENVIRONMENTS):
        for pid in FOCUS_PERSONAS_BY_ENV[env_idx]:
            p = pmap[pid]
            exp = expectation(env, p)
            focus_rows.append((
                f"{env[0]}_{pid}", env[1], f"{pid} {p['name']}", exp[0], exp[1], exp[2], exp[3], exp[4], exp[5], exp[6], exp[7]
            ))
    add_table(
        doc,
        ["CASE_ID", "环境", "人格", "预期感知", "预期情绪", "触发记忆", "激活信念", "需求缺失", "形成欲望", "预期行为", "不能出现的错误路径"],
        focus_rows,
        [850, 850, 850, 1000, 820, 1350, 1200, 800, 1200, 1200, 1240],
        font_size=5,
    )

    doc.add_heading("七、批量矩阵测试：2000 个组合", level=1)
    doc.add_paragraph("矩阵用于自动化测试和评分。每一行都绑定环境、人格、预期行为方向、核心心理原因、是否允许多种答案，以及禁止出现的模板化错误。")
    matrix_rows = []
    for env in ENVIRONMENTS:
        for p in personas:
            d, reason = direction(env, p)
            forbidden = "禁止只输出“谨慎观察/保持边界/逐步信任”，必须绑定该人格的记忆、信念或需求。"
            matrix_rows.append((f"{env[0]}_{p['pid']}", env[0], p["pid"], d, reason, "允许，只要心理链路一致", forbidden))
    add_table(
        doc,
        ["CASE_ID", "ENV", "PID", "预期行为方向", "核心心理原因", "多答案", "不能出现的错误路径"],
        matrix_rows,
        [1050, 700, 700, 1200, 2300, 1400, 2010],
        font_size=5,
    )

    doc.add_heading("八、同结果异因测试", level=1)
    same_behavior_rows = [
        ("拒绝加入高风险项目", "被背叛者", "怕被骗，陌生合作关系触发背叛记忆。"),
        ("拒绝加入高风险项目", "长期贫困者", "怕失去最后一点安全垫，贫困记忆压过翻身欲望。"),
        ("拒绝加入高风险项目", "完美主义者", "信息不完整，害怕失败证明自己不够好。"),
        ("拒绝加入高风险项目", "责任型人格", "不能让家人或团队承担自己的风险选择。"),
        ("拒绝加入高风险项目", "原则主义者", "项目价值观或规则边界不干净。"),
        ("拒绝加入高风险项目", "回避依恋", "不想被陌生团队绑定或索取情感承诺。"),
        ("拒绝加入高风险项目", "悲观主义者", "默认高回报背后是失败概率。"),
        ("拒绝加入高风险项目", "控制型人格", "自己没有主导权，无法掌控关键变量。"),
    ]
    add_table(doc, ["相同行为", "人格", "必须区分的心理原因"], same_behavior_rows, [1800, 1800, 5760], font_size=8)

    doc.add_heading("九、人格连续性测试", level=1)
    continuity_rows = [
        ("PID_021 被背叛者", "ENV_001 高风险合作", "警惕陌生人动机，要求合同、资金流和退出机制。"),
        ("PID_021 被背叛者", "ENV_002 朋友借钱", "不直接拒绝，但核验用途、还款能力和书面约定。"),
        ("PID_021 被背叛者", "ENV_006 亲密冷淡", "过度解读信号，担心对方隐瞒真实意图。"),
        ("PID_021 被背叛者", "ENV_007 权威命令", "怀疑权威把风险转嫁给自己，要求责任边界。"),
        ("PID_021 被背叛者", "成长修复状态", "能觉察自己过度防御，允许低风险试探但不跳过核验。"),
    ]
    add_table(doc, ["人格", "环境", "连续性预期"], continuity_rows, [1800, 2000, 5560], font_size=8)

    doc.add_heading("十、成长漂移测试", level=1)
    drift_rows = [
        ("旧模式触发", "角色按旧信念解释事件，例如“别人接近我就是利用我”。"),
        ("新证据进入", "环境提供稳定、可复验、低风险的正向证据。"),
        ("信念松动", "角色不应瞬间改变，而是出现“也许并非所有靠近都是利用”的条件性修正。"),
        ("行为微调", "从拒绝变成低风险试行，从回避变成有限沟通，从攻击变成设边界。"),
        ("边界保留", "成长不是失去防御，而是防御变得可解释、可调节、可被证据更新。"),
    ]
    add_table(doc, ["阶段", "判定标准"], drift_rows, [1800, 7560], font_size=8)

    doc.add_heading("十一、评分标准", level=1)
    scoring_rows = [
        ("行为合理性", "0-10", "行为在当前人格和环境下是否合理。"),
        ("心理链路完整性", "0-10", "是否完整输出感知、情绪、记忆、信念、需求、欲望、行为。"),
        ("人格区分度", "0-10", "同一环境下，不同人格是否走出不同路径。"),
        ("人格连续性", "0-10", "同一人格在不同环境下是否保持核心一致。"),
        ("创伤触发准确性", "0-10", "是否正确触发相关经历，而不是乱触发。"),
        ("成长漂移合理性", "0-10", "经历新事件后，人格变化是否合理、不突兀。"),
        ("反模板化程度", "0-10", "是否避免所有结果都变成“谨慎、观察、边界表达”。"),
    ]
    add_table(doc, ["评分项", "分值", "说明"], scoring_rows, [2200, 900, 6260], font_size=8)

    doc.add_heading("十二、基础心理链路测试 V1：旧版 100 Case 保留", level=1)
    doc.add_paragraph("以下为原 CharacterOS 角色行为 100 个测试用例，作为基础心理链路回归集。V2 不删除旧版，而是把它定位为单人格单场景心理合理性测试。")
    add_table(doc, old_rows[0], old_rows[1:], [1550, 1450, 1900, 1600, 1800, 1060], font_size=6, header_fill="F2F4F7")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "CharacterOS Contradiction & Differentiation Test Suite V2"

    doc.core_properties.title = "CharacterOS 人格分化与矛盾测试集 V2"
    doc.core_properties.subject = "CharacterOS benchmark v2"
    doc.core_properties.author = "Codex"
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
